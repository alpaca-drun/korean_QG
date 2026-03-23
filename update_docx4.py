import docx
from docx.shared import Pt
from docx.oxml.ns import qn
import copy

try:
    doc = docx.Document('app/download/sample3.docx')
    for table in doc.tables:
        found = False
        for row in table.rows:
            for cell in row.cells:
                if '{answer_explain}' in cell.text:
                    found = True
                    break
            if found: break
            
        if found:
            # Helper to copy row
            def copy_row(source_row):
                new_row = copy.deepcopy(source_row._tr)
                source_row._tr.addnext(new_row)
                return docx.table._Row(new_row, table)
            
            last_row = table.rows[-1]
            fmt_run_label = last_row.cells[0].paragraphs[0].runs[0] if last_row.cells[0].paragraphs[0].runs else None
            fmt_run_value = last_row.cells[-1].paragraphs[0].runs[0] if last_row.cells[-1].paragraphs[0].runs else None
            
            def set_cell_text_with_fmt(cell, text, ref_run):
                cell.text = ''
                if ref_run:
                    new_run = cell.paragraphs[0].add_run(text)
                    new_run.font.name = '맑은 고딕'
                    new_run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    new_run.font.size = Pt(10)
                    new_run.font.color.rgb = ref_run.font.color.rgb
                else:
                    new_run = cell.paragraphs[0].add_run(text)
                    new_run.font.name = '맑은 고딕'
                    new_run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                    new_run.font.size = Pt(10)

            new_row1 = copy_row(last_row)
            set_cell_text_with_fmt(new_row1.cells[0], '인정답안', fmt_run_label)
            set_cell_text_with_fmt(new_row1.cells[-1], '{accepted_answers}', fmt_run_value)
            
            new_row2 = copy_row(new_row1)
            set_cell_text_with_fmt(new_row2.cells[0], '채점기준', fmt_run_label)
            set_cell_text_with_fmt(new_row2.cells[-1], '{scoring_criteria}', fmt_run_value)
            
            doc.save('app/download/sample3.docx')
            print('Successfully updated sample3.docx')
            break
except Exception as e:
    print('Error:', str(e))
