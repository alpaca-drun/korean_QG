from docx import Document
from copy import deepcopy
import os
import re
from dotenv import load_dotenv
import sys
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import json
from app.db.database import select_one

# 실행 위치에 상관없이 import 되도록 경로 보정
CURRENT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.abspath(os.path.join(CURRENT_DIR, "..", ".."))
if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)

from app.db.database import get_db_connection
from app.core.logger import logger


def execute_query_via_app_db(query: str, params: tuple | None = None, fetch: bool = True):
    """
    app/db/storage.py의 get_db_connection()을 사용해서 쿼리를 실행합니다.
    (FastAPI 서버와 동일한 settings/db 환경변수를 사용)
    """
    try:
        with get_db_connection() as connection:
            with connection.cursor() as cursor:
                cursor.execute(query, params)
                if fetch:
                    return cursor.fetchall()
                return cursor.rowcount
    except Exception as e:
        logger.error(f"DB 쿼리 실행 실패: {e}")
        raise e

# .env 파일에서 환경변수 로드
load_dotenv()
def get_cell_text(table, row, col):
    try:
        cell = table.cell(row, col)
        text = cell.text.strip()
        if not text and col > 0:
            text = table.cell(row, col - 1).text.strip()
        return text
    except IndexError:
        return ""
def find_table_in_cell(cell, tag, doc=None):
    """
    셀 안에 있는 표를 재귀적으로 찾는 함수
    
    Args:
        cell: 셀 객체
        tag: 찾을 태그 문자열
        doc: Document 객체 (표 객체 생성 시 필요)
    
    Returns:
        찾은 Table 객체 또는 None
    """
    from docx.table import Table
    
    # 셀 안의 모든 표 요소 찾기 (XPath 사용)
    try:
        # lxml의 xpath 사용
        tbl_elements = cell._element.xpath('.//w:tbl', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
    except:
        # xpath가 없는 경우 직접 찾기
        tbl_elements = []
        for elem in cell._element.iter():
            if elem.tag.endswith('}tbl'):
                tbl_elements.append(elem)
    
    for tbl_elm in tbl_elements:
        try:
            # 표 객체 생성 (doc이 필요함)
            if doc is not None:
                nested_table = Table(tbl_elm, doc)
            else:
                # doc이 없으면 직접 접근 시도
                nested_table = Table(tbl_elm, cell._parent._parent)
            
            # 표 안의 모든 셀 검색
            for row in nested_table.rows:
                for nested_cell in row.cells:
                    cell_text = nested_cell.text
                    if tag in cell_text:
                        logger.info("중첩 표를 찾았습니다! (셀 안의 표)")
                        logger.debug("찾은 셀 내용: %s", cell_text[:100])
                        return nested_table
            
            # 중첩 표 안에 또 다른 표가 있을 수 있으므로 재귀적으로 검색
            for row in nested_table.rows:
                for nested_cell in row.cells:
                    result = find_table_in_cell(nested_cell, tag, doc)
                    if result:
                        return result
        except Exception as e:
            logger.debug("표 객체 생성/검색 실패 (다음 표로): %s", e)
            continue
    
    return None

def find_career_table(doc, tag="{answer}"):
    """
    표에서 특정 태그를 포함한 표를 찾는 함수 (중첩 표 포함)
    
    Args:
        doc: Document 객체
        tag: 찾을 태그 문자열 (기본값: "{answer}")
    
    Returns:
        찾은 Table 객체 또는 None
    """
    # 최상위 레벨의 표 검색
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                # 셀의 텍스트 직접 가져오기
                cell_text = cell.text
                if tag in cell_text:
                    logger.info("표를 찾았습니다! (표 인덱스: %s, 행: %s, 열: %s)", table_idx, row_idx, col_idx)
                    logger.debug("찾은 셀 내용: %s", cell_text[:100])
                    return table
                
                # 셀 안에 중첩된 표가 있는지 확인
                nested_table = find_table_in_cell(cell, tag, doc)
                if nested_table:
                    return nested_table
    
    logger.warning("'%s' 태그를 포함한 표를 찾을 수 없습니다.", tag)
    # 디버깅: 모든 표의 첫 번째 셀 내용 로깅
    logger.debug("디버깅 정보 - 모든 표의 첫 번째 셀 내용:")
    for table_idx, table in enumerate(doc.tables):
        if len(table.rows) > 0 and len(table.rows[0].cells) > 0:
            first_cell_text = table.rows[0].cells[0].text[:50]
            logger.debug("표 %s: %s...", table_idx, first_cell_text)
    return None
def extract_category_from_info_id(info_id):
    """
    CREATE_INFO_ID에서 카테고리 값을 추출하는 함수
    예: "말하기듣기_30-05-05" -> "말하기듣기"
    
    Args:
        info_id: 정보 ID 문자열
    
    Returns:
        카테고리 문자열 (없으면 빈 문자열)
    """
    if not info_id:
        logger.debug("[카테고리 추출] info_id가 없습니다.")
        return ""
    
    # 언더스코어로 분리하여 첫 번째 부분 추출
    parts = str(info_id).split('_')
    if len(parts) > 0:
        category = parts[0]
        logger.debug("[카테고리 추출] '%s' → '%s'", info_id, category)
        return category
    logger.debug("[카테고리 추출] '%s'에서 카테고리를 추출할 수 없습니다.", info_id)
    return ""

def replace_document_text(doc, replacements):
    """
    문서 전체에서 플레이스홀더를 교체하는 함수 (표 외부의 텍스트 포함)
    
    Args:
        doc: Document 객체
        replacements: 플레이스홀더와 값의 딕셔너리 (예: {'{category}': '말하기듣기'})
    """
    logger.info("[문서 플레이스홀더 교체] 시작 (교체할 항목: %s개)", len(replacements))
    replaced_count = 0
    
    # 문서의 모든 단락에서 교체
    for paragraph in doc.paragraphs:
        if paragraph.text:
            new_text = paragraph.text
            for placeholder, value in replacements.items():
                if placeholder in new_text:
                    new_text = new_text.replace(placeholder, value)
                    replaced_count += 1
            
            if new_text != paragraph.text:
                # 기존 run의 서식 정보 저장
                reference_run = None
                if paragraph.runs:
                    reference_run = paragraph.runs[0]
                
                # 단락 내용 교체
                paragraph.clear()
                if new_text:
                    new_run = paragraph.add_run(new_text)
                    if reference_run:
                        copy_run_formatting(reference_run, new_run)
    
    # 표 안의 셀에서도 교체 (표 내부는 replace_table_text에서 처리되지만, 
    # 표 외부의 플레이스홀더를 위해 여기서도 처리)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text:
                        new_text = paragraph.text
                        for placeholder, value in replacements.items():
                            if placeholder in new_text:
                                new_text = new_text.replace(placeholder, value)
                                replaced_count += 1
                        
                        if new_text != paragraph.text:
                            # 기존 run의 서식 정보 저장
                            reference_run = None
                            if paragraph.runs:
                                reference_run = paragraph.runs[0]

                            paragraph.clear()
                            if new_text:
                                new_run = paragraph.add_run(new_text)
                                if reference_run:
                                    copy_run_formatting(reference_run, new_run)
    
    logger.info("[문서 플레이스홀더 교체] 완료 (총 %s개 교체)", replaced_count)

def fill_table_from_list(doc_path, output_path, data_list, category=""):
    """
    sample.docx의 표를 복사하여 리스트 데이터로 채우는 함수
    
    Args:
        doc_path: 원본 docx 파일 경로
        output_path: 출력 파일 경로
        data_list: 표를 채울 데이터 리스트 (각 항목은 dict 형태)
                  예: [{'nu': 1, 'question': '질문1', 'select1': '선택1', ...}, ...]
        category: 카테고리 문자열 (예: "말하기듣기", "쓰기", "매체")
    """
    logger.info("=" * 60)
    logger.info("[문서 처리 시작] 입력: %s, 출력: %s, 데이터: %s개, 카테고리: %s", doc_path, output_path, len(data_list), category or "(없음)")
    logger.info("=" * 60)
    
    # 원본 문서 열기
    logger.info("[1/5] 문서 열기 중...")
    doc = Document(doc_path)
    logger.info("문서 열기 완료 (표 개수: %s개)", len(doc.tables))
    
    # 카테고리 플레이스홀더 교체 (문서 전체)
    if category:
        logger.info("[2/5] 카테고리 플레이스홀더 교체 중...")
        replace_document_text(doc, {'{category}': category})
    else:
        logger.info("[2/5] 카테고리 플레이스홀더 교체 건너뜀 (카테고리 없음)")
    
    # 첫 번째 표 찾기 (원본 표)
    logger.info("[3/5] 표 찾기 중...")
    if len(doc.tables) == 0:
        logger.error("표를 찾을 수 없습니다!")
        return
    
    # {answer} 플레이스홀더가 포함된 표 찾기
    original_table = find_career_table(doc, "{answer}")
    if original_table is None:
        logger.error("{answer} 태그가 포함된 표를 찾을 수 없습니다.")
        return
    
    logger.info("원본 표 찾기 완료")

    # 원본 표의 element를 저장 (플레이스홀더가 있는 원본 상태를 먼저 저장)
    logger.info("[4/5] 표 데이터 채우기 중...")
    original_table_elm = deepcopy(original_table._element)
    previous_table_elm = original_table._element
    
    num = 1
    # 첫 번째 데이터로 원본 표 채우기
    if data_list:
        logger.info("표 %s/%s 채우는 중...", num, len(data_list))
        replace_table_text(original_table, data_list[0], num)
        logger.debug("표 %s/%s 완료", num, len(data_list))
        
        # 나머지 데이터에 대해 표 복사 및 채우기
        for data in data_list[1:]:
            num += 1 
            logger.info("표 %s/%s 채우는 중...", num, len(data_list))
            
            # 원본 표 element 복사 (플레이스홀더가 있는 원본 상태로 복사)
            new_table_elm = deepcopy(original_table_elm)
            
            # 이전 표 다음에 줄바꿈(단락) 추가
            from docx.oxml import OxmlElement
            
            # 단락 요소 생성 (빈 줄)
            p = OxmlElement('w:p')
            previous_table_elm.addnext(p)
            
            # 단락 다음에 새 표 삽입
            p.addnext(new_table_elm)
            
            # 새로 추가된 표를 Document 객체로 찾기
            new_table_obj = None
            for t in doc.tables:
                if t._element == new_table_elm:
                    new_table_obj = t
                    break
            
            if new_table_obj:
                replace_table_text(new_table_obj, data, num)
                previous_table_elm = new_table_elm
            logger.debug("표 %s/%s 완료", num, len(data_list))
    
    # 결과 저장
    logger.info("[5/5] 파일 저장 중...")
    doc.save(output_path)
    logger.info("저장 완료!")
    logger.info("=" * 60)
    logger.info("완료! %s개의 표가 생성되어 %s에 저장되었습니다.", len(data_list), output_path)
    logger.info("=" * 60)

def get_project_id_from_env_or_arg(project_id: str | int | None = None) -> int:
    """
    프로젝트 ID를 환경변수/인자에서 안전하게 가져옵니다.

    우선순위:
    - 인자 project_id
    - PROJECT_ID
    - CREATE_PROJECT_ID
    - CREATE_INFO_ID (기존 호환: 숫자면 project_id로 간주)
    """
    if project_id is None:
        project_id = (
            os.getenv("PROJECT_ID")
            or os.getenv("CREATE_PROJECT_ID")
            or os.getenv("CREATE_INFO_ID")
        )

    row = select_one(
        "projects", 
        where={"project_id": project_id}, 
        columns="project_id, project_name" 
    )

    project_name = row.get("project_name") if row else f"project-{project_id}"

    if not project_id:
        raise ValueError("PROJECT_ID 환경변수가 설정되지 않았습니다. (또는 CREATE_PROJECT_ID/CREATE_INFO_ID 숫자값)")
    try:
        return int(str(project_id).strip()), project_name
    except ValueError:
        raise ValueError(f"PROJECT_ID가 정수가 아닙니다: {project_id}")


def get_project_passage_text(project_id: int, user_id: int | None = None) -> str:
    """
    project_source_config를 기반으로 프로젝트의 지문(원본/커스텀)을 가져옵니다.
    - custom_passage_id가 있으면 passage_custom.context
    - passage_id가 있으면 passages.context
    """
    query = """
        SELECT
            psc.custom_passage_id,
            psc.passage_id,
            pc.context AS custom_context,
            ps.context AS passage_context
        FROM project_source_config psc
        LEFT JOIN passage_custom pc ON pc.custom_passage_id = psc.custom_passage_id
        LEFT JOIN passages ps ON ps.passage_id = psc.passage_id
        WHERE psc.project_id = %s
        ORDER BY psc.created_at DESC
        LIMIT 1
    """
    # 프로젝트 소유권 검증(선택)
    if user_id is not None:
        ownership = execute_query_via_app_db(
            "SELECT project_id FROM projects WHERE project_id = %s AND user_id = %s AND is_deleted = 0 LIMIT 1",
            params=(project_id, user_id),
            fetch=True,
        )
        if not ownership:
            return ""

    results = execute_query_via_app_db(query, params=(project_id,), fetch=True)
    if not results:
        return ""
    row = results[0] or {}
    return (row.get("custom_context") or row.get("passage_context") or "").strip()



def get_matching_question_data(project_id: int | None = None, user_id: int | None = None):
    """
    선긋기 전용 데이터를 DB에서 조회하여 반환하는 함수 (정답 포맷팅 포함)
    """
    project_id_int = project_id
    
    query = """
        SELECT
            mq.matching_question_id AS qid,
            mq.created_at AS created_at,
            mq.question AS question,
            NULLIF(mq.modified_passage, '') AS passage,
            NULL AS answer,
            mq.answer_explain AS answer_explain,
            NULL AS box_content,
            4 AS qtype,
            mq.left_items AS left_items,
            mq.right_items AS right_items,
            mq.sort_order AS sort_order
        FROM matching_questions mq
        JOIN projects p ON p.project_id = mq.project_id
        WHERE mq.project_id = %s
    """
    
    try:
        if user_id is None:
            base_filters = " AND 1=1 AND p.is_deleted = 0"
            params = (project_id_int,)
        else:
            base_filters = " AND p.user_id = %s AND p.is_deleted = 0"
            params = (project_id_int, user_id)

        filtered_query = query.replace("WHERE mq.project_id = %s", f"WHERE mq.project_id = %s{base_filters} AND IFNULL(mq.is_checked, 0) = 1")
        filtered_query += " ORDER BY qid ASC"
        
        results = execute_query_via_app_db(filtered_query, params=params, fetch=True)
        
        if not results:
            logger.warning("project_id=%s에 해당하는 선긋기 문항 데이터가 없습니다.", project_id_int)
            return []
            
        data_list = []
        left_symbols = ['Ⓐ', 'Ⓑ', 'Ⓒ', 'Ⓓ', 'Ⓔ']
        right_symbols = ['①', '②', '③', '④', '⑤']

        for idx, row in enumerate(results, 1):
            selects = {}
            extra_data = {}
            formatted_answer = ""
            
            try:
                left_items = json.loads(row.get('left_items') or '[]')
                right_items = json.loads(row.get('right_items') or '[]')
                sort_order = row.get('sort_order')
                
                if isinstance(sort_order, str):
                    try:
                        sort_order = json.loads(sort_order)
                    except:
                        sort_order = []
                
                display_rights = []
                # sort_order 유효성 검사 및 정렬
                if sort_order and isinstance(sort_order, list) and len(sort_order) == len(right_items):
                     display_rights = [right_items[i] for i in sort_order]
                else:
                     display_rights = right_items
                     sort_order = list(range(len(right_items)))
                
                # 데이터 채우기 (left1~5, right1~5)
                for i in range(5):
                    l_item = left_items[i] if i < len(left_items) else ""
                    r_item = display_rights[i] if i < len(display_rights) else ""
                    
                    extra_data[f'left{i+1}'] = l_item
                    extra_data[f'right{i+1}'] = r_item
                    
                    selects[f'select{i+1}'] = f"{l_item}   ----------------   {r_item}" if l_item and r_item else ""

                # 정답 포맷팅 (Ⓐ-②, Ⓑ-①)
                answer_parts = []
                for i in range(len(left_items)):
                    # 왼쪽 i번째 항목의 짝은 right_items[i]임.
                    # right_items[i]가 화면의 몇 번째(k)에 있는지 찾아야 함.
                    # 즉, sort_order[k] == i 인 k를 찾아야 함.
                    try:
                        if i in sort_order:
                            k = sort_order.index(i)
                            # 기호 매핑 (범위 체크)
                            l_sym = left_symbols[i] if i < len(left_symbols) else f"L{i+1}"
                            r_sym = right_symbols[k] if k < len(right_symbols) else f"R{k+1}"
                            answer_parts.append(f"{l_sym}-{r_sym}")
                    except ValueError:
                        pass
                
                if answer_parts:
                    formatted_answer = ", ".join(answer_parts)
                else:
                    formatted_answer = row.get('answer', '')

            except Exception as e:
                logger.error(f"선긋기 데이터 파싱 오류: {e}")
                formatted_answer = row.get('answer', '')

            item = {
                'qid': row.get('qid'),
                'nu': idx,
                'question': row.get('question', ''),
                'select1': selects.get('select1') or '',
                'select2': selects.get('select2') or '',
                'select3': selects.get('select3') or '',
                'select4': selects.get('select4') or '',
                'select5': selects.get('select5') or '',
                'answer': formatted_answer,
                'answer_explain': row.get('answer_explain', ''),
                'passage': row.get('passage', ''),
                'boxcontent': row.get('box_content', '')
            }
            item.update(extra_data)
            data_list.append(item)
            
        return data_list

    except Exception as e:
        logger.exception("[DB 연결 오류] 예상치 못한 오류: %s", e)
        raise


def get_question_data_from_db(project_id: int | None = None, user_id: int | None = None):
    """
    DB에서 질문(객관식/단답형/OX) 데이터를 가져오는 함수
    
    Args:
        project_id: 프로젝트 ID
    
    Returns:
        질문 데이터 리스트 (각 항목은 dict 형태)
    """
    # project_id_int = get_project_id_from_env_or_arg(project_id)
    project_id_int = project_id
    # passage_text = get_project_passage_text(project_id_int, user_id=user_id)
    # logger.debug("passage_text: %s", passage_text)
    
    # ✅ 현재 DB 스키마 기반: multiple_choice_questions / short_answer_questions / true_false_questions / matching_questions
    # seq는 생성시간 기준으로 부여
    query = """
        (
            SELECT
                mcq.question_id AS qid,
                mcq.created_at AS created_at,
                mcq.question AS question,
                NULLIF(mcq.modified_passage, '') AS passage,
                mcq.option1 AS select1,
                mcq.option2 AS select2,
                mcq.option3 AS select3,
                mcq.option4 AS select4,
                mcq.option5 AS select5,
                mcq.answer AS answer,
                mcq.answer_explain AS answer_explain,
                mcq.box_content AS box_content,
                1 AS qtype,
                NULL AS left_items,
                NULL AS right_items,
                NULL AS sort_order,
                NULL AS accepted_answers,
                NULL AS scoring_criteria
            FROM multiple_choice_questions mcq
            JOIN projects p ON p.project_id = mcq.project_id
            WHERE mcq.project_id = %s
        )
        UNION ALL
        (
            SELECT
                saq.short_question_id AS qid,
                saq.created_at AS created_at,
                saq.question AS question,
                NULLIF(saq.modified_passage, '') AS passage,
                NULL AS select1,
                NULL AS select2,
                NULL AS select3,
                NULL AS select4,
                NULL AS select5,
                saq.answer AS answer,
                saq.answer_explain AS answer_explain,
                saq.box_content AS box_content,
                2 AS qtype,
                NULL AS left_items,
                NULL AS right_items,
                NULL AS sort_order,
                NULL AS accepted_answers,
                NULL AS scoring_criteria
            FROM short_answer_questions saq
            JOIN projects p2 ON p2.project_id = saq.project_id
            WHERE saq.project_id = %s
        )
        UNION ALL
        (
            SELECT
                tfq.ox_question_id AS qid,
                tfq.created_at AS created_at,
                tfq.question AS question,
                NULLIF(tfq.modified_passage, '') AS passage,
                'O' AS select1,
                'X' AS select2,
                NULL AS select3,
                NULL AS select4,
                NULL AS select5,
                tfq.answer AS answer,
                tfq.answer_explain AS answer_explain,
                NULL AS box_content,
                3 AS qtype,
                NULL AS left_items,
                NULL AS right_items,
                NULL AS sort_order,
                NULL AS accepted_answers,
                NULL AS scoring_criteria
            FROM true_false_questions tfq
            JOIN projects p3 ON p3.project_id = tfq.project_id
            WHERE tfq.project_id = %s
        )
        UNION ALL
        (
            SELECT
                mq.matching_question_id AS qid,
                mq.created_at AS created_at,
                mq.question AS question,
                NULLIF(mq.modified_passage, '') AS passage,
                NULL AS select1,
                NULL AS select2,
                NULL AS select3,
                NULL AS select4,
                NULL AS select5,
                NULL AS answer,
                mq.answer_explain AS answer_explain,
                NULL AS box_content,
                4 AS qtype,
                mq.left_items AS left_items,
                mq.right_items AS right_items,
                mq.sort_order AS sort_order,
                NULL AS accepted_answers,
                NULL AS scoring_criteria
            FROM matching_questions mq
            JOIN projects p4 ON p4.project_id = mq.project_id
            WHERE mq.project_id = %s
        )
        UNION ALL
        (
            SELECT
                laq.long_question_id AS qid,
                laq.created_at AS created_at,
                laq.question AS question,
                NULLIF(laq.modified_passage, '') AS passage,
                NULL AS select1,
                NULL AS select2,
                NULL AS select3,
                NULL AS select4,
                NULL AS select5,
                laq.answer AS answer,
                laq.answer_explain AS answer_explain,
                laq.box_content AS box_content,
                5 AS qtype,
                NULL AS left_items,
                NULL AS right_items,
                NULL AS sort_order,
                laq.accepted_answers AS accepted_answers,
                laq.scoring_criteria AS scoring_criteria
            FROM long_answer_questions laq
            JOIN projects p5 ON p5.project_id = laq.project_id
            WHERE laq.project_id = %s
        )
        ORDER BY qid ASC
    """
    
    # DB 연결 설정 확인 (try 블록 밖에서 정의하여 except에서도 사용 가능)
    # env_prefix = os.getenv('DB_ENV_PREFIX', 'QG_db')
    # database = os.getenv('DB_DATABASE', 'midtest')
    
    try:
        
        # print(f"🔌 [DB 연결] project_id={project_id_int}로 데이터 조회 중...")
        # print(f"   환경변수 접두사: {env_prefix}")
        # print(f"   데이터베이스: {database}")
        
        # # 환경변수 확인
        # host = os.getenv(f'{env_prefix}_host')
        # user = os.getenv(f'{env_prefix}_user')
        # password = os.getenv(f'{env_prefix}_password')
        
        # if not host:
        #     print(f"   ⚠️ 경고: {env_prefix}_host 환경변수가 설정되지 않았습니다.")
        # if not user:
        #     print(f"   ⚠️ 경고: {env_prefix}_user 환경변수가 설정되지 않았습니다.")
        # if not password:
        #     print(f"   ⚠️ 경고: {env_prefix}_password 환경변수가 설정되지 않았습니다.")
        
        # print(f"   DB 연결 시도 중...")
        # 프로젝트 소유권/삭제 여부 필터링(선택)
        if user_id is None:
            base_filters = " AND 1=1 AND p.is_deleted = 0"
            params = (project_id_int, project_id_int, project_id_int, project_id_int, project_id_int)
        else:
            base_filters = " AND p.user_id = %s AND p.is_deleted = 0"
            # p2/p3/p4/p5도 동일하게 적용되도록 문자열 치환
            params = (project_id_int, user_id, project_id_int, user_id, project_id_int, user_id, project_id_int, user_id, project_id_int, user_id)

        filtered_query = (
            query
            .replace("WHERE mcq.project_id = %s", f"WHERE mcq.project_id = %s{base_filters} AND IFNULL(mcq.is_checked, 0) = 1")
            .replace(
                "WHERE saq.project_id = %s",
                (f"WHERE saq.project_id = %s AND p2.user_id = %s AND p2.is_deleted = 0 AND IFNULL(saq.is_checked, 0) = 1")
                if user_id is not None
                else "WHERE saq.project_id = %s AND IFNULL(saq.is_checked, 0) = 1"
            )
            .replace(
                "WHERE tfq.project_id = %s",
                (f"WHERE tfq.project_id = %s AND p3.user_id = %s AND p3.is_deleted = 0 AND IFNULL(tfq.is_checked, 0) = 1")
                if user_id is not None
                else "WHERE tfq.project_id = %s AND IFNULL(tfq.is_checked, 0) = 1"
            )
            .replace(
                "WHERE mq.project_id = %s",
                (f"WHERE mq.project_id = %s AND p4.user_id = %s AND p4.is_deleted = 0 AND IFNULL(mq.is_checked, 0) = 1")
                if user_id is not None
                else "WHERE mq.project_id = %s AND IFNULL(mq.is_checked, 0) = 1"
            )
            .replace(
                "WHERE laq.project_id = %s",
                (f"WHERE laq.project_id = %s AND p5.user_id = %s AND p5.is_deleted = 0 AND IFNULL(laq.is_checked, 0) = 1")
                if user_id is not None
                else "WHERE laq.project_id = %s AND IFNULL(laq.is_checked, 0) = 1"
            )
        )

        results = execute_query_via_app_db(filtered_query, params=params, fetch=True)
        
        if not results:
            logger.warning("project_id=%s에 해당하는 문항 데이터가 없습니다.", project_id_int)
            return []
        
        logger.info("DB 쿼리 완료 (조회된 행: %s개)", len(results))
        
        # 결과를 딕셔너리 리스트로 변환
        logger.info("[데이터 변환] 딕셔너리로 변환 중...")
        data_list = []
        for idx, row in enumerate(results, 1):
            
            # 선긋기(qtype=4) 처리
            qtype = row.get('qtype')
            selects = {}
            extra_data = {}

            if qtype == 4:
                try:
                    left_items = json.loads(row.get('left_items') or '[]')
                    right_items = json.loads(row.get('right_items') or '[]')
                    sort_order = row.get('sort_order')
                    
                    if isinstance(sort_order, str):
                        try:
                            sort_order = json.loads(sort_order)
                        except:
                            sort_order = []
                    
                    display_rights = []
                    if sort_order and isinstance(sort_order, list) and len(sort_order) == len(right_items):
                         display_rights = [right_items[i] for i in sort_order]
                    else:
                         display_rights = right_items
                    
                    for i in range(5):
                        l_item = left_items[i] if i < len(left_items) else ""
                        r_item = display_rights[i] if i < len(display_rights) else ""
                        
                        extra_data[f'left{i+1}'] = l_item
                        extra_data[f'right{i+1}'] = r_item
                        
                        selects[f'select{i+1}'] = f"{l_item}   ----------------   {r_item}" if l_item and r_item else ""
                except Exception as e:
                    logger.error(f"선긋기 데이터 파싱 오류: {e}")

            # 번호는 전체 문항 순서로 부여
            item = {
                'qid': row.get('qid'),
                'nu': idx,
                'question': row.get('question', ''),
                'select1': selects.get('select1') or row.get('select1', '') or '',
                'select2': selects.get('select2') or row.get('select2', '') or '',
                'select3': selects.get('select3') or row.get('select3', '') or '',
                'select4': selects.get('select4') or row.get('select4', '') or '',
                'select5': selects.get('select5') or row.get('select5', '') or '',
                'answer': row.get('answer', ''),
                'answer_explain': row.get('answer_explain', ''),
                'passage': row.get('passage', ''),
                'boxcontent': row.get('box_content', ''),
                'accepted_answers': row.get('accepted_answers', ''),
                'scoring_criteria': row.get('scoring_criteria', '')
            }
            
            if qtype == 4:
                item.update(extra_data)
            
            data_list.append(item)
            if idx % 10 == 0 or idx == len(results):
                logger.debug("진행 중... %s/%s", idx, len(results))
        
        logger.info("변환 완료! 총 %s개의 질문 데이터를 가져왔습니다.", len(data_list))
        return data_list
        
    except ValueError as e:
        logger.error("[DB 연결 오류] 설정 오류 발생: %s", e)
        logger.info("해결 방법: .env 또는 환경변수에 DB_HOST, DB_PORT, DB_USER, DB_PASSWORD, DB_DATABASE 설정")
        raise
    except Exception as e:
        logger.exception("[DB 연결 오류] 예상치 못한 오류: %s - %s", type(e).__name__, e)
        logger.info("해결 방법: DB 서버 실행 여부, 네트워크, 방화벽, 환경변수 확인")
        raise

def copy_run_formatting(source_run, target_run):
    """
    source_run의 서식(XML rPr)을 target_run으로 전체 복사하여
    한글/영문 폰트, 크기, 스타일을 완벽하게 보존함
    """
    try:
        # 원본 Run의 서식 XML(rPr)을 가져옴
        source_rPr = source_run._element.rPr
        if source_rPr is not None:
            # 타겟 Run의 기존 rPr 제거
            target_rPr = target_run._element.rPr
            if target_rPr is not None:
                target_run._element.remove(target_rPr)
            
            # 원본 rPr을 복사하여 타겟 Run의 첫 번째 자식으로 삽입
            target_run._element.insert(0, deepcopy(source_rPr))
            
    except Exception as e:
        logger.debug("서식 복사 중 오류: %s", e)
        pass

def parse_markdown_table_data(table_lines):
    """마크다운 표 라인을 파싱하여 2차원 리스트로 반환"""
    data = []
    for i, line in enumerate(table_lines):
        # 구분선(---|---)은 건너뜀
        if i == 1 and re.match(r'^\s*\|?[\s\-:|]+\|?\s*$', line):
            continue
        
        # 셀 분리 (양끝 | 제거 후 split)
        row_content = line.strip().strip('|')
        cells = [c.strip() for c in row_content.split('|')]
        data.append(cells)
    return data

def parse_markdown_text(text):
    """텍스트에서 마크다운 표를 감지하여 텍스트와 표 데이터로 분리"""
    # \n 리터럴을 실제 줄바꿈 문자로 변환
    if text:
        text = text.replace('\\n', '\n')
        
    segments = []
    lines = text.split('\n')
    current_text = []
    table_lines = []
    in_table = False
    
    for line in lines:
        stripped = line.strip()
        # 표 감지 로직
        if stripped.startswith('|') and stripped.endswith('|'):
            if not in_table:
                if current_text:
                    segments.append({'type': 'text', 'content': '\n'.join(current_text)})
                    current_text = []
                in_table = True
            table_lines.append(stripped)
        else:
            if in_table:
                if table_lines:
                    # 유효한 표인지 확인 (2줄 이상, 두 번째 줄이 구분선 패턴)
                    if len(table_lines) >= 2 and re.match(r'^\s*\|?[\s\-:|]+\|?\s*$', table_lines[1]):
                         segments.append({'type': 'table', 'content': parse_markdown_table_data(table_lines)})
                    else:
                         segments.append({'type': 'text', 'content': '\n'.join(table_lines)})
                    table_lines = []
                in_table = False
            current_text.append(line)
            
    # 잔여 처리
    if in_table and table_lines:
        if len(table_lines) >= 2 and re.match(r'^\s*\|?[\s\-:|]+\|?\s*$', table_lines[1]):
             segments.append({'type': 'table', 'content': parse_markdown_table_data(table_lines)})
        else:
             segments.append({'type': 'text', 'content': '\n'.join(table_lines)})
    elif current_text:
        segments.append({'type': 'text', 'content': '\n'.join(current_text)})
        
    return segments

def apply_inline_styles(paragraph, text, base_run=None):
    """
    텍스트 내의 인라인 스타일(<u>, **)을 파싱하여 paragraph에 run으로 추가
    (줄바꿈 문자 \n 처리 포함)
    """
    if not text:
        return
        
    # 줄바꿈 문자로 먼저 분리
    lines = text.split('\n')
    
    # <u>...</u> 또는 **...** 패턴 찾기 (그룹핑으로 분리)
    pattern = r'(<u>.*?</u>|\*\*.*?\*\*)'
    
    for i, line in enumerate(lines):
        # 첫 번째 줄이 아니면 줄바꿈(Shift+Enter 효과) 추가
        if i > 0:
            paragraph.add_run().add_break()
            
        if not line:
            continue

        parts = re.split(pattern, line)
        
        for part in parts:
            if not part:
                continue
                
            run_text = part
            is_underline = False
            is_bold = False
            
            # 태그 확인 및 제거
            if part.startswith('<u>') and part.endswith('</u>'):
                run_text = part[3:-4]
                is_underline = True
            elif part.startswith('**') and part.endswith('**'):
                run_text = part[2:-2]
                is_bold = True
                
            if not run_text:
                continue
                
            # Run 추가
            new_run = paragraph.add_run(run_text)
            
            # 기본 서식 복사
            if base_run:
                copy_run_formatting(base_run, new_run)
                
            # 스타일 적용
            if is_underline:
                new_run.font.underline = True
            if is_bold:
                new_run.font.bold = True

def insert_markdown_content(cell, paragraph, markdown_segments, base_run=None):
    """셀 내의 특정 단락 뒤에 마크다운 세그먼트들을 삽입"""
    current_element = paragraph._element
    
    # 전달받은 base_run이 없으면 paragraph의 첫 번째 run 사용
    if base_run is None and paragraph.runs:
        base_run = paragraph.runs[0]
        
    for segment in markdown_segments:
        if segment['type'] == 'text':
            content = segment['content']
            if not content.strip():
                continue
            
            try:
                # 새 단락 생성
                temp_p = cell.add_paragraph() 
                
                # 1. 단락 스타일(Style ID) 복사
                if paragraph.style:
                    temp_p.style = paragraph.style
                
                # 2. 단락 속성(pPr - 정렬, 줄간격 등) 복사
                if paragraph._element.pPr is not None:
                    if temp_p._element.pPr is not None:
                        temp_p._element.remove(temp_p._element.pPr)
                    temp_p._element.insert(0, deepcopy(paragraph._element.pPr))
                
                # 인라인 스타일 적용하여 텍스트 추가
                apply_inline_styles(temp_p, content, base_run)
                
                temp_elm = temp_p._element
                temp_elm.getparent().remove(temp_elm)
                current_element.addnext(temp_elm)
                current_element = temp_elm
            except Exception as e:
                logger.error(f"마크다운 텍스트 삽입 실패: {e}")

        elif segment['type'] == 'table':
            table_data = segment['content']
            if not table_data:
                continue
            
            try:
                rows = len(table_data)
                cols = max(len(row) for row in table_data) if rows > 0 else 0
                
                if rows > 0 and cols > 0:
                    temp_table = cell.add_table(rows=rows, cols=cols)
                    temp_table.style = 'Table Grid'
                    
                    for r, row_data in enumerate(table_data):
                        for c, cell_text in enumerate(row_data):
                            if c < cols:
                                cell_obj = temp_table.cell(r, c)
                                cell_obj.text = cell_text
                                # 폰트 스타일 설정 (필요시 추가)
                                # for p in cell_obj.paragraphs:
                                #     if p.runs and base_run:
                                #         copy_run_formatting(base_run, p.runs[0])

                    tbl_elm = temp_table._element
                    tbl_elm.getparent().remove(tbl_elm)
                    current_element.addnext(tbl_elm)
                    current_element = tbl_elm
                    
                    # 표 뒤에 빈 줄 추가
                    spacer_p = OxmlElement('w:p')
                    current_element.addnext(spacer_p)
                    current_element = spacer_p
            except Exception as e:
                logger.error(f"마크다운 표 삽입 실패: {e}")

def replace_table_text(table, data, num):
    """
    표의 플레이스홀더를 실제 데이터로 교체하는 함수 (서식 유지)
    
    Args:
        table: docx Table 객체
        data: 채울 데이터 (dict)
        num: 문항 번호
    """
    # 플레이스홀더 교체 딕셔너리
    # data.get(key) 가 None일 경우 ''로 처리하여 문자열 "None"이 생성되는 것 방지
    replacements = {
        '{num}': str(num),
        '{question}': str(data.get('question') or ''),
        '{select1}': str(data.get('select1') or ''),
        '{select2}': str(data.get('select2') or ''),
        '{select3}': str(data.get('select3') or ''),
        '{select4}': str(data.get('select4') or ''),
        '{select5}': str(data.get('select5') or ''),
        '{left1}': str(data.get('left1') or ''),
        '{left2}': str(data.get('left2') or ''),
        '{left3}': str(data.get('left3') or ''),
        '{left4}': str(data.get('left4') or ''),
        '{left5}': str(data.get('left5') or ''),
        '{right1}': str(data.get('right1') or ''),
        '{right2}': str(data.get('right2') or ''),
        '{right3}': str(data.get('right3') or ''),
        '{right4}': str(data.get('right4') or ''),
        '{right5}': str(data.get('right5') or ''),
        '{answer}': str(data.get('answer') or ''),
        '{answer_explain}': str(data.get('answer_explain') or ''),
        '{passage}': str(data.get('passage') or ''),
        '{boxcontent}': str(data.get('boxcontent') or ''),
        '{accepted_answers}': str(data.get('accepted_answers') or ''),
        '{scoring_criteria}': str(data.get('scoring_criteria') or '')
    }
    
    # 1. 값이 비어있는 경우 해당 행 삭제 처리
    rows_to_delete = []
    
    # 1-1. 일반 플레이스홀더 목록 ({passage} 제외)
    # 이 목록에 있는 플레이스홀더는 값이 비어있으면 해당 행을 삭제함
    check_placeholders = [
        '{question}', '{select1}', '{select2}', '{select3}', '{select4}', '{select5}', 
        '{left1}', '{left2}', '{left3}', '{left4}', '{left5}',
        '{right1}', '{right2}', '{right3}', '{right4}', '{right5}',
        '{answer}', '{answer_explain}', '{boxcontent}',
        '{accepted_answers}', '{scoring_criteria}'
    ]
    
    for row in table.rows:
        row_text = "".join(cell.text for cell in row.cells)
        should_delete_row = False
        
        # 일반 플레이스홀더 체크
        for placeholder in check_placeholders:
            if placeholder in row_text:
                value = replacements.get(placeholder, '')
                # 값이 없거나, 빈 문자열이거나, '-' 이거나, 문자열 "None"인 경우 행 삭제
                if not value or str(value).strip() == '' or str(value).strip() == '-' or str(value).strip().lower() == 'none':
                    should_delete_row = True
                    break
        
        if should_delete_row:
             rows_to_delete.append(row)
             continue 

        # 1-2. {passage} 별도 체크
        # {passage}는 마크다운 표 등이 들어올 수 있으므로 값이 있을 때는 삭제하지 않음
        # 값이 없을 때만 삭제
        if '{passage}' in row_text:
             val = replacements.get('{passage}', '')
             # 디버깅: {passage} 값 로깅
             logger.debug(f"[DEBUG] passage 행 확인: 값='{val}'")
             
             if not val or str(val).strip() == '' or str(val).strip() == '-' or str(val).strip().lower() == 'none':
                 logger.debug("[DEBUG] passage 행 삭제 대상 포함됨 (값이 비어있음)")
                 rows_to_delete.append(row)
            
    # 행 제거 (뒤에서부터 삭제하여 인덱스 꼬임 방지)
            
    # 행 제거 (뒤에서부터 삭제하여 인덱스 꼬임 방지)
    # 삭제 시 인접한 행들의 테두리도 정리
    for row in reversed(rows_to_delete):
        # 현재 행의 인덱스 찾기
        try:
            current_idx = -1
            for i, r in enumerate(table.rows):
                if r._tr == row._tr:
                    current_idx = i
                    break
            
            if current_idx != -1:
                # 1. 이전 행의 하단 테두리 제거
                if current_idx > 0:
                    prev_row = table.rows[current_idx - 1]
                    for cell in prev_row.cells:
                        _set_cell_border(cell, bottom={"val": "nil"})
                
                # 2. 다음 행의 상단 테두리 제거
                if current_idx < len(table.rows) - 1:
                    next_row = table.rows[current_idx + 1]
                    for cell in next_row.cells:
                        _set_cell_border(cell, top={"val": "nil"})
        except:
            pass # 인덱스 조회 실패 시 건너뜀

        tr = row._tr
        parent = tr.getparent()
        if parent is not None:
            parent.remove(tr)

    # 2. 남은 표 내의 모든 셀을 순회하며 플레이스홀더 교체
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            
            # 2-1. 마크다운/인라인 스타일이 포함된 플레이스홀더 처리
            processed_paragraphs = [] # 이미 처리된 단락 추적
            
            # 셀 내의 모든 단락을 순회하며 처리
            # list()로 복사하여 순회 중 수정에 대비
            for p in list(cell.paragraphs):
                p_text = p.text
                if not p_text: continue
                
                # 해당 단락에 어떤 플레이스홀더가 있는지 확인
                found_placeholders = []
                for ph, val in replacements.items():
                    if ph in p_text:
                        found_placeholders.append((ph, val))
                
                if not found_placeholders:
                    continue
                
                # 발견된 플레이스홀더 중 마크다운/스타일 처리가 필요한 것 확인
                needs_special_processing = False
                target_ph = None
                target_val = None

                for ph, val in found_placeholders:
                    # 마크다운 표(|...|) 또는 인라인 스타일(<u>, **)이 있는지 확인
                    # 표는 최소 2줄 이상이어야 하므로 newline 체크
                    has_markdown_table = ('|' in val and '\n' in val)
                    has_inline_style = ('<u>' in val or '**' in val)
                    
                    if has_markdown_table or has_inline_style:
                        needs_special_processing = True
                        target_ph = ph
                        target_val = val
                        break
                
                if needs_special_processing and target_ph:
                    # 1. 서식(폰트, 크기 등) 보존을 위해 원본 Run 객체 확보
                    # 플레이스홀더가 포함된 run을 우선적으로 찾음
                    base_run = None
                    if p.runs:
                        for run in p.runs:
                            if target_ph in run.text:
                                base_run = run
                                break
                        # 없으면 첫 번째 run 사용
                        if base_run is None:
                            base_run = p.runs[0]
                    
                    # 2. 플레이스홀더 제거 (단락 내 텍스트 치환)
                    # p.text = ""를 하면 p.runs도 모두 사라질 수 있으므로 base_run을 미리 확보해야 함
                    if p.text.strip() == target_ph:
                            p.text = "" 
                    else:
                            p.text = p.text.replace(target_ph, "")
                    
                    # 3. 마크다운 파싱 및 삽입 (base_run 전달)
                    segments = parse_markdown_text(target_val)
                    insert_markdown_content(cell, p, segments, base_run=base_run)
                    
                    processed_paragraphs.append(p)
                    # 한 단락에 여러 특수 처리가 필요한 경우 복잡해질 수 있으므로
                    # 일단 하나 처리하면 해당 단락 처리는 완료된 것으로 간주 (단순화)
                    continue
            
            # 2-2. 일반 텍스트 치환 (기존 로직)
            # 마크다운 처리가 안 된 단락들만 대상
            for paragraph in cell.paragraphs:
                if paragraph in processed_paragraphs:
                    continue
                # 단락의 전체 텍스트 확인
                para_text = paragraph.text
                if not para_text:
                    continue
                
                # 플레이스홀더가 있는지 확인
                has_placeholder = False
                for placeholder in replacements.keys():
                    if placeholder in para_text:
                        has_placeholder = True
                        break
                
                if not has_placeholder:
                    continue
                
                # 플레이스홀더를 실제 값으로 교체
                replaced_text = para_text
                for placeholder, value in replacements.items():
                    if placeholder in replaced_text:
                        replaced_text = replaced_text.replace(placeholder, value)
                
                # 텍스트가 변경되었는지 확인
                if replaced_text == para_text:
                    continue
                
                # 기존 run들의 서식 정보 저장 (첫 번째 run의 서식 사용)
                reference_run = None
                if paragraph.runs:
                    reference_run = paragraph.runs[0]
                
                # 기존 run들을 모두 제거
                for run in list(paragraph.runs):
                    run._element.getparent().remove(run._element)
                
                # 교체된 텍스트를 새 run으로 추가 (서식 유지)
                if replaced_text:
                    new_run = paragraph.add_run(replaced_text)
                    if reference_run:
                        copy_run_formatting(reference_run, new_run)

def _set_cell_border(cell, **kwargs):
    """
    셀의 테두리를 설정하는 내부 유틸리티 함수
    Usage: _set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"})
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))

# 사용 예시
if __name__ == "__main__":
    import sys
    
    logger.info("=" * 60)
    logger.info("문서 생성 스크립트 시작")
    logger.info("=" * 60)
    
    # 프로젝트 ID 결정
    project_id = os.getenv("PROJECT_ID") or os.getenv("CREATE_PROJECT_ID") or os.getenv("CREATE_INFO_ID")
    logger.info("[환경변수 확인] PROJECT_ID/CREATE_PROJECT_ID/CREATE_INFO_ID = %s", project_id or "(설정되지 않음)")
    if not project_id:
        logger.error("PROJECT_ID 환경변수가 설정되지 않았습니다. (또는 CREATE_PROJECT_ID/CREATE_INFO_ID 숫자값)")
        sys.exit(1)

    project_id_int, project_name = get_project_id_from_env_or_arg(project_id)
    category = os.getenv("CATEGORY", "")
    
    # DB에서 데이터 가져오기
    try:
        logger.info("[DB 데이터 조회] 시작...")
        data_list = get_question_data_from_db(project_id_int)
        
        if not data_list:
            logger.error("가져올 데이터가 없습니다.")
            sys.exit(1)
        
        # 입력 파일과 출력 파일 경로 (환경변수에서 가져오거나 기본값 사용)
        input_file = os.getenv('INPUT_DOCX', 'sample3.docx')
        output_file = os.getenv('OUTPUT_DOCX', f'{project_name}.docx')
        
        logger.info("[파일 경로] 입력: %s, 출력: %s", input_file, output_file)
        
        # 함수 실행 (카테고리 전달)
        fill_table_from_list(input_file, output_file, data_list, category=category)
        
    except ValueError as e:
        logger.error("오류: %s", e)
        logger.info("사용법: 환경변수 설정 export INFO_ID=123 또는 명령줄 인자 python dev.py 123")
        sys.exit(1)
    except Exception as e:
        logger.exception("예상치 못한 오류 발생: %s", e)
        sys.exit(1)
    
    # 기존 샘플 데이터 (사용 안 함)
    _sample_data = [
        {
            'nu': 1,
            'question': '첫 번째 질문입니다',
            'select1': '선택지 1-1',
            'select2': '선택지 1-2',
            'select3': '선택지 1-3',
            'select4': '선택지 1-4',
            'select5': '선택지 1-5'
        },
        {
            'nu': 2,
            'question': '두 번째 질문입니다',
            'select1': '선택지 2-1',
            'select2': '선택지 2-2',
            'select3': '선택지 2-3',
            'select4': '선택지 2-4',
            'select5': '선택지 2-5'
        },
        {
            'nu': 3,
            'question': '세 번째 질문입니다',
            'select1': '선택지 3-1',
            'select2': '선택지 3-2',
            'select3': '선택지 3-3',
            'select4': '선택지 3-4',
            'select5': '선택지 3-5'
        },
        {
            'nu': 3,
            'question': '세 번째 질문입니다',
            'select1': '선택지 3-1',
            'select2': '선택지 3-2',
            'select3': '선택지 3-3',
            'select4': '선택지 3-4',
            'select5': '선택지 3-5'
        },
        {
            'nu': 3,
            'question': '세 번째 질문입니다',
            'select1': '선택지 3-1',
            'select2': '선택지 3-2',
            'select3': '선택지 3-3',
            'select4': '선택지 3-4',
            'select5': '선택지 3-5'
        },
        {
            'nu': 3,
            'question': '세 번째 질문입니다',
            'select1': '선택지 3-1',
            'select2': '선택지 3-2',
            'select3': '선택지 3-3',
            'select4': '선택지 3-4',
            'select5': '선택지 3-5'
        },
        {
            'nu': 3,
            'question': '세 번째 질문입니다',
            'select1': '선택지 3-1',
            'select2': '선택지 3-2',
            'select3': '선택지 3-3',
            'select4': '선택지 3-4',
            'select5': '선택지 3-5'
        },
        {
            'nu': 3,
            'question': '세 번째 질문입니다',
            'select1': '선택지 3-1',
            'select2': '선택지 3-2',
            'select3': '선택지 3-3',
            'select4': '선택지 3-4',
            'select5': '선택지 3-5'
        },
        {
            'nu': 3,
            'question': '세 번째 질문입니다',
            'select1': '선택지 3-1',
            'select2': '선택지 3-2',
            'select3': '선택지 3-3',
            'select4': '선택지 3-4',
            'select5': '선택지 3-5'
        },
        {
            'nu': 3,
            'question': '세 번째 질문입니다',
            'select1': '선택지 3-1',
            'select2': '선택지 3-2',
            'select3': '선택지 3-3',
            'select4': '선택지 3-4',
            'select5': '선택지 3-5'
        },
        {
            'nu': 3,
            'question': '세 번째 질문입니다',
            'select1': '선택지 3-1',
            'select2': '선택지 3-2',
            'select3': '선택지 3-3',
            'select4': '선택지 3-4',
            'select5': '선택지 3-5'
        },
        {
            'nu': 3,
            'question': '세 번째 질문입니다',
            'select1': '선택지 3-1',
            'select2': '선택지 3-2',
            'select3': '선택지 3-3',
            'select4': '선택지 3-4',
            'select5': '선택지 3-5'
        },
        {
            'nu': 3,
            'question': '세 번째 질문입니다',
            'select1': '선택지 3-1',
            'select2': '선택지 3-2',
            'select3': '선택지 3-3',
            'select4': '선택지 3-4',
            'select5': '선택지 3-5'
        }
    ]