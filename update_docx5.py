import docx
import copy
from docx.oxml import parse_xml

try:
    doc = docx.Document('app/download/sample3.docx')
    for table in doc.tables:
        found = False
        for row in table.rows:
            for cell in row.cells:
                if '{answer_explain}' in cell.text:
                    found = True
                    break
            if found: break
            
        if found:
            # Helper to copy row
            def copy_row(source_row):
                new_row = copy.deepcopy(source_row._tr)
                source_row._tr.addnext(new_row)
                return docx.table._Row(new_row, table)
            
            last_row = table.rows[-1]
            
            def set_cell_text_with_fmt(cell, text):
                cell.text = ''
                new_run = cell.paragraphs[0].add_run(text)
                rPr_xml = '''
                <w:rPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                    <w:rFonts w:ascii="맑은 고딕" w:hAnsi="맑은 고딕" w:eastAsia="맑은 고딕" w:cs="맑은 고딕"/>
                    <w:sz w:val="20"/>
                    <w:szCs w:val="20"/>
                </w:rPr>
                '''
                rPr = parse_xml(rPr_xml)
                new_run._element.append(rPr)

            new_row1 = copy_row(last_row)
            set_cell_text_with_fmt(new_row1.cells[0], '인정답안')
            set_cell_text_with_fmt(new_row1.cells[-1], '{accepted_answers}')
            
            new_row2 = copy_row(new_row1)
            set_cell_text_with_fmt(new_row2.cells[0], '채점기준')
            set_cell_text_with_fmt(new_row2.cells[-1], '{scoring_criteria}')
            
            doc.save('app/download/sample3.docx')
            print('Successfully updated sample3.docx with 맑은 고딕 10pt')
            break
except Exception as e:
    print('Error:', str(e))
