import docx
import copy

try:
    doc = docx.Document('app/download/sample3.docx')
    for table in doc.tables:
        found = False
        for row in table.rows:
            for cell in row.cells:
                if '{answer_explain}' in cell.text:
                    found = True
                    break
            if found: break
            
        if found:
            # Helper to copy row
            def copy_row(source_row):
                new_row = copy.deepcopy(source_row._tr)
                source_row._tr.addnext(new_row)
                return docx.table._Row(new_row, table)
            
            # The last row is '해설' (cells 0, 1, 2) and '{answer_explain}' (cell 3)
            # Add row for scoring_criteria first so it appears before accepted_answers,
            # or vice versa. Let's do accepted_answers then scoring_criteria.
            
            new_row1 = copy_row(table.rows[-1])
            # Modify the runs to preserve formatting
            for run in new_row1.cells[0].paragraphs[0].runs:
                if '해설' in run.text:
                    run.text = run.text.replace('해설', '인정답안')
            for run in new_row1.cells[3].paragraphs[0].runs:
                if '{answer_explain}' in run.text:
                    run.text = run.text.replace('{answer_explain}', '{accepted_answers}')
            
            new_row2 = copy_row(new_row1)
            for run in new_row2.cells[0].paragraphs[0].runs:
                if '인정답안' in run.text:
                    run.text = run.text.replace('인정답안', '채점기준')
            for run in new_row2.cells[3].paragraphs[0].runs:
                if '{accepted_answers}' in run.text:
                    run.text = run.text.replace('{accepted_answers}', '{scoring_criteria}')
            
            doc.save('app/download/sample3.docx')
            print('Successfully updated sample3.docx')
            break
except Exception as e:
    print('Error:', str(e))
